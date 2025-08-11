# utils/doc_utils.py
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_COLOR_INDEX
import os

def read_docx_text(path: str):
    """Return full text and list of paragraph strings."""
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paragraphs)
    return full_text, paragraphs

def save_reviewed_docx(original_path: str, annotations: list, output_path: str):
    """
    Insert visible annotation paragraphs after flagged paragraphs.
    annotations is list of dicts with keys:
      paragraph_index, flag_text, comment, severity, suggestion, citation
    """
    doc = Document(original_path)
    new_doc = Document()
    p_index = 0

    for p in doc.paragraphs:
        new_p = new_doc.add_paragraph()
        for run in p.runs:
            r = new_p.add_run(run.text)
            try:
                r.bold = run.bold
                r.italic = run.italic
                r.underline = run.underline
                r.font.size = run.font.size
            except Exception:
                pass

        matched = [a for a in annotations if a.get("paragraph_index") == p_index]
        for a in matched:
            ann_p = new_doc.add_paragraph()
            ann_run = ann_p.add_run(f"[REVIEW NOTE - {a.get('severity','Medium')}] {a.get('comment')}")
            ann_run.italic = True
            ann_run.font.size = Pt(10)
            ann_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

            if a.get("suggestion"):
                s_p = new_doc.add_paragraph()
                s_run = s_p.add_run(f"Suggestion: {a.get('suggestion')}")
                s_run.font.size = Pt(10)

            if a.get("citation"):
                c_p = new_doc.add_paragraph()
                c_run = c_p.add_run(f"Citation / Source excerpt:\n{a.get('citation')}")
                c_run.font.size = Pt(9)
        p_index += 1

    outdir = os.path.dirname(output_path) or "."
    os.makedirs(outdir, exist_ok=True)
    try:
        new_doc.save(output_path)
    except Exception as e:
        # attempt to backup existing file if conflict
        if os.path.exists(output_path) and not os.path.isdir(output_path):
            os.rename(output_path, output_path + "_backup")
            new_doc.save(output_path)
        else:
            raise e
    return output_path
